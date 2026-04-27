from __future__ import annotations

from io import BytesIO
from pathlib import Path
import base64

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
import subprocess
import tempfile
import sys
from fastapi.concurrency import run_in_threadpool
from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.models import Booking
from app.settings_service import get_merged_settings


BASE_DIR = Path(__file__).resolve().parent.parent
UPLOADS_DIR = BASE_DIR / "uploads"
TEMPLATES_DIR = Path(__file__).resolve().parent / "templates"

jinja_env = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


def _format_inr(value: float | int | None) -> str:
    return f"\u20b9{float(value or 0):,.0f}".replace(",", "_").replace("_", ",")


def _build_summary(booking: Booking) -> str:
    source = (booking.ai_summary or "").strip()
    if source:
        return source

    location = _selected_location_label(booking)
    budget = _format_inr(booking.budget) if booking.budget else "the planned budget"
    description = " ".join((booking.polished_description or booking.ad_description or "brand promotion campaign").split())
    return (
        f"This booking is scheduled for {location} with a projected spend of {budget}, "
        f"designed to promote {description} with focused visibility and timely audience reach."
    )


def _status_meta(status: str | None) -> dict[str, str]:
    normalized = (status or "pending").lower()
    if normalized in {"confirmed", "approved"}:
      return {"label": "Approved", "class_name": "status-approved"}
    if normalized in {"cancelled", "rejected"}:
      return {"label": "Rejected", "class_name": "status-rejected"}
    if normalized == "completed":
      return {"label": "Completed", "class_name": "status-completed"}
    return {"label": "Pending", "class_name": "status-pending"}


def _billing_label(value: str | None) -> str:
    return value or "Custom Package"


def _image_to_data_uri(url: str | None) -> str:
    if not url:
        return ""

    if url.startswith("data:"):
        return url

    if url.startswith("/uploads/"):
        filename = url.replace("/uploads/", "", 1)
        file_path = UPLOADS_DIR / filename
        if file_path.exists():
            mime = "image/png"
            suffix = file_path.suffix.lower()
            if suffix in {".jpg", ".jpeg"}:
                mime = "image/jpeg"
            elif suffix == ".webp":
                mime = "image/webp"
            encoded = base64.b64encode(file_path.read_bytes()).decode("ascii")
            return f"data:{mime};base64,{encoded}"

    return url


def _selected_location_snapshots(booking: Booking) -> list[dict]:
    snapshots = booking.selected_screens or []
    if snapshots:
        return snapshots
    if booking.screen:
        return [
            {
                "id": booking.screen.id,
                "name": booking.screen.name,
                "area": booking.screen.area,
            }
        ]
    return []


def _selected_location_label(booking: Booking) -> str:
    snapshots = _selected_location_snapshots(booking)
    if not snapshots:
        return "the selected location"

    names = [snapshot.get("name") or f'Location #{snapshot.get("id", "N/A")}' for snapshot in snapshots]
    if len(names) <= 2:
        return ", ".join(names)
    return f"{names[0]}, {names[1]} + {len(names) - 2} more"


def _data_uri_to_bytes(data_uri: str | None) -> bytes | None:
    if not data_uri or not data_uri.startswith("data:"):
        return None
    try:
        encoded = data_uri.split(",", 1)[1]
        return base64.b64decode(encoded)
    except Exception:
        return None


def build_invoice_context(booking: Booking) -> dict:
    merged = get_merged_settings()
    config = merged.get("config", {})
    prefix = (config.get("invoice_prefix") or "INV").strip() or "INV"
    safe_prefix = "".join(ch for ch in prefix if ch.isalnum() or ch in {"-", "_"})
    company_name = config.get("invoice_company_name") or config.get("general_app_name") or "OLRAC Advertising"
    contact_phone = config.get("general_contact_phone") or merged.get("whatsapp_number") or ""
    status = _status_meta(booking.status.value if booking.status else "pending")

    return {
        "company_name": company_name,
        "logo_url": _image_to_data_uri(config.get("invoice_logo_url") or config.get("general_logo_url")),
        "seal_image": _image_to_data_uri(config.get("invoice_seal_url")),
        "contact_email": merged.get("contact_email") or "",
        "contact_phone": contact_phone,
        "address": config.get("invoice_address") or "",
        "gst_number": config.get("invoice_gst") or "",
        "invoice_id": f"{safe_prefix}-{str(booking.id).zfill(5)}",
        "invoice_date": booking.created_at.strftime("%d %B %Y") if booking.created_at else "",
        "client_name": booking.client_name or "N/A",
        "client_company": booking.company or "Not provided",
        "client_phone": booking.phone or "Not provided",
        "location": _selected_location_label(booking),
        "booking_date": booking.created_at.strftime("%d %B %Y") if booking.created_at else "",
        "time_slot": _billing_label(booking.duration_label),
        "status_label": status["label"],
        "status_class": status["class_name"],
        "quantity": booking.slot_quantity or 1,
        "amount": _format_inr(booking.total_price),
        "total": _format_inr(booking.total_price),
        "slot_quantity": booking.slot_quantity or 1,
        "summary": _build_summary(booking),
        "invoice_title": config.get("invoice_title") or "INVOICE",
        "primary_color": config.get("invoice_primary_color") or "#334155",
        "show_seal": bool(config.get("invoice_show_seal", True)),
        "show_slot_details": bool(config.get("invoice_show_slot_details", True)),
        "signature_name": config.get("invoice_signature_name") or company_name,
        "signature_title": config.get("invoice_signature_title") or "Authorized Signatory",
        "payment_terms": config.get("invoice_payment_terms") or "",
        "notes": config.get("invoice_notes") or "",
        "footer_note": config.get("invoice_footer_note") or "",
        "bank_name": config.get("invoice_bank_name") or "",
        "account_name": config.get("invoice_account_name") or "",
        "account_number": config.get("invoice_account_number") or "",
        "ifsc": config.get("invoice_ifsc") or "",
        "upi": config.get("invoice_upi") or "",
    }


def build_quotation_context(booking: Booking) -> dict:
    merged = get_merged_settings()
    config = merged.get("config", {})
    prefix = (config.get("quotation_prefix") or "QTN").strip() or "QTN"
    safe_prefix = "".join(ch for ch in prefix if ch.isalnum() or ch in {"-", "_"})
    company_name = config.get("quotation_company_name") or config.get("general_app_name") or "OLRAC Advertising"
    contact_phone = config.get("general_contact_phone") or merged.get("whatsapp_number") or ""

    return {
        "company_name": company_name,
        "logo_url": _image_to_data_uri(config.get("quotation_logo_url") or config.get("general_logo_url")),
        "seal_url": _image_to_data_uri(config.get("quotation_seal_url") or config.get("invoice_seal_url")),
        "contact_email": merged.get("contact_email") or "",
        "contact_phone": contact_phone,
        "address": config.get("quotation_address") or "",
        "gst_number": config.get("quotation_gst") or "",
        "quotation_id": f"{safe_prefix}-{str(booking.id).zfill(5)}",
        "location": _selected_location_label(booking),
        "primary_color": config.get("quotation_primary_color") or "#334155",
        "bank_name": config.get("quotation_bank_name") or "",
        "account_name": config.get("quotation_account_name") or "",
        "account_number": config.get("quotation_account_number") or "",
        "ifsc": config.get("quotation_ifsc") or "",
        "upi": config.get("quotation_upi") or "",
        "tax_rate": config.get("quotation_tax_rate") or "18",
    }


def render_invoice_html(context: dict) -> str:
    template = jinja_env.get_template("invoice.html")
    return template.render(**context)


def _run_subprocess_sync(script: str, html: bytes) -> tuple[bytes, bytes, int]:
    proc = subprocess.run(
        [sys.executable, "-c", script],
        input=html,
        capture_output=True
    )
    return proc.stdout, proc.stderr, proc.returncode

async def _playwright_pdf(html_content: str, *, margin: str = "14mm", viewport_width: int = 1280) -> bytes:
    with tempfile.NamedTemporaryFile("wb", suffix=".pdf", delete=False) as f:
        pdf_path = f.name

    pdf_out = pdf_path.replace("\\", "/")

    script = f"""
from playwright.sync_api import sync_playwright
import sys

html_content = sys.stdin.buffer.read().decode('utf-8')

try:
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={{"width": {viewport_width}, "height": 1123}})
        page.set_content(html_content, wait_until="domcontentloaded")
        page.pdf(path='{pdf_out}', format="A4", print_background=True, margin={{"top": "{margin}", "right": "{margin}", "bottom": "{margin}", "left": "{margin}"}})
        browser.close()
except Exception as e:
    import traceback
    traceback.print_exc()
    sys.exit(1)
"""

    import os
    stdout, stderr, returncode = await run_in_threadpool(_run_subprocess_sync, script, html_content.encode('utf-8'))

    if returncode != 0:
        error_msg = stderr.decode('utf-8', errors='replace') if stderr else stdout.decode('utf-8', errors='replace')
        if os.path.exists(pdf_path):
            try: os.remove(pdf_path)
            except: pass
        raise Exception(f"Playwright Subprocess Failed:\\n{error_msg}")

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    if os.path.exists(pdf_path):
        try: os.remove(pdf_path)
        except: pass

    return pdf_bytes


async def generate_invoice_pdf_bytes(context: dict) -> bytes:
    return await _playwright_pdf(render_invoice_html(context))


async def generate_pdf_from_html(html_content: str) -> bytes:
    """Render an A4 document HTML (794px body) to PDF via Playwright."""
    return await _playwright_pdf(html_content, margin="0mm", viewport_width=794)


def _set_run_style(run, *, size=10.5, bold=False, color="0F172A"):
    run.font.name = "Segoe UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_borders(cell, color="D7DEE8", size=6, value="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def _clear_cell_borders(cell):
    _set_cell_borders(cell, color="FFFFFF", size=0, value="nil")


def _set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_mm: float):
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(width_mm * 56.7)))
    width.set(qn("w:type"), "dxa")


def _set_table_fixed_width(table, widths_mm: list[float]):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for index, width in enumerate(widths_mm):
            _set_cell_width(row.cells[index], width)


def _clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)


def _shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_paragraph(cell, text="", *, bold=False, size=10.5, color="0F172A", align=None, spacing_after=4):
    paragraph = cell.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    _set_run_style(run, size=size, bold=bold, color=color)
    return paragraph


def _add_body_paragraph(document, text="", *, bold=False, size=10.5, color="0F172A", align=None, spacing_after=4):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    _set_run_style(run, size=size, bold=bold, color=color)
    return paragraph


def _status_fill(status_class: str) -> tuple[str, str]:
    mapping = {
        "status-approved": ("DCFCE7", "166534"),
        "status-rejected": ("FFE4E6", "BE185D"),
        "status-completed": ("E0F2FE", "0369A1"),
        "status-pending": ("FFEDD5", "C2410C"),
    }
    return mapping.get(status_class, ("FFEDD5", "C2410C"))


def generate_invoice_docx_bytes(context: dict) -> bytes:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10)

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    logo_bytes = _data_uri_to_bytes(context.get("logo_url"))
    seal_bytes = _data_uri_to_bytes(context.get("seal_image"))

    header = document.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_width(header, [118, 64])
    _clear_table_borders(header)
    left_cell, right_cell = header.rows[0].cells
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(left_cell, top=80, start=0, bottom=80, end=120)
    _set_cell_margins(right_cell, top=80, start=120, bottom=80, end=0)

    brand_table = left_cell.add_table(rows=1, cols=2)
    _set_table_fixed_width(brand_table, [24, 88])
    _clear_table_borders(brand_table)
    brand_logo, brand_name = brand_table.rows[0].cells
    brand_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    brand_name.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(brand_logo, top=0, start=0, bottom=0, end=60)
    _set_cell_margins(brand_name, top=0, start=60, bottom=0, end=0)
    if logo_bytes:
        p = brand_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(BytesIO(logo_bytes), width=Mm(18))
    brand_name_p = brand_name.paragraphs[0]
    brand_name_p.paragraph_format.space_after = Pt(0)
    brand_name_run = brand_name_p.add_run(context["company_name"])
    _set_run_style(brand_name_run, size=16, bold=True)

    invoice_title = right_cell.paragraphs[0]
    invoice_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    invoice_title.paragraph_format.space_after = Pt(2)
    run = invoice_title.add_run(context.get("invoice_title") or "INVOICE")
    _set_run_style(run, size=22, bold=True)
    _add_paragraph(right_cell, f'Invoice ID: {context["invoice_id"]}', size=9, color="475569", align=WD_ALIGN_PARAGRAPH.RIGHT, spacing_after=1)
    _add_paragraph(right_cell, f'Date: {context["invoice_date"]}', size=9, color="475569", align=WD_ALIGN_PARAGRAPH.RIGHT, spacing_after=0)

    divider = document.add_paragraph()
    divider.paragraph_format.space_before = Pt(14)
    divider.paragraph_format.space_after = Pt(12)
    divider_run = divider.add_run(" " * 250)
    divider_run.font.underline = True
    divider_run.font.color.rgb = RGBColor.from_string("D7DEE8")

    details = document.add_table(rows=1, cols=2)
    details.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_width(details, [86, 86])
    _clear_table_borders(details)
    client_cell, booking_cell = details.rows[0].cells
    _set_cell_margins(client_cell, top=0, start=0, bottom=40, end=240)
    _set_cell_margins(booking_cell, top=0, start=240, bottom=40, end=0)
    _add_paragraph(client_cell, "Client Details", bold=True, size=12, spacing_after=8)
    for label, value in [
        ("Name", context["client_name"]),
        ("Company", context["client_company"]),
        ("Phone", context["client_phone"]),
    ]:
        p = client_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_label = p.add_run(f"{label}  ")
        _set_run_style(run_label, size=10.5, bold=True, color="64748B")
        run_value = p.add_run(str(value))
        _set_run_style(run_value, size=10.5)

    _add_paragraph(booking_cell, "Booking Details", bold=True, size=12, spacing_after=8)
    for label, value in [
        ("Location", context["location"]),
        ("Date", context["booking_date"]),
        ("Time Slot", context["time_slot"]),
        *(
            [
                ("Slots Booked", str(context.get("slot_quantity") or 1)),
            ]
            if context.get("show_slot_details")
            else []
        ),
    ]:
        if not value:
            continue
        p = booking_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_label = p.add_run(f"{label}  ")
        _set_run_style(run_label, size=10.5, bold=True, color="64748B")
        run_value = p.add_run(str(value))
        _set_run_style(run_value, size=10.5)

    status_p = booking_cell.add_paragraph()
    status_p.paragraph_format.space_after = Pt(0)
    run_label = status_p.add_run("Status  ")
    _set_run_style(run_label, size=10.5, bold=True, color="64748B")
    fill, color = _status_fill(context["status_class"])
    badge_run = status_p.add_run(context["status_label"])
    _set_run_style(badge_run, size=9.5, bold=True, color=color)

    _add_body_paragraph(document, "Pricing Table", bold=True, size=12, spacing_after=8)
    pricing = document.add_table(rows=2, cols=4)
    pricing.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_width(pricing, [94, 28, 30, 30])
    headers = ["Description", "Quantity", "Amount", "Total"]
    for index, title in enumerate(headers):
        cell = pricing.rows[0].cells[index]
        _shade_cell(cell, "1F2937")
        _set_cell_borders(cell, color="1F2937")
        _set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        _set_run_style(run, size=9.5, bold=True, color="FFFFFF")
    description = "Ad Slot Booking"
    if context.get("show_slot_details"):
        description = f'{description} ({context.get("slot_quantity") or 1} slots, {context["time_slot"]})'
    values = [description, str(context["quantity"]), context["amount"], context["total"]]
    for index, value in enumerate(values):
        cell = pricing.rows[1].cells[index]
        _set_cell_borders(cell)
        _set_cell_margins(cell, top=160, start=120, bottom=160, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if index >= 2 else WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        _set_run_style(run, size=10.5, bold=index == 3)

    total_table = document.add_table(rows=1, cols=2)
    total_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_fixed_width(total_table, [44, 52])
    total_left, total_right = total_table.rows[0].cells
    _shade_cell(total_left, "F8FAFC")
    _shade_cell(total_right, "F8FAFC")
    _set_cell_borders(total_left, color="F8FAFC")
    _set_cell_borders(total_right, color="F8FAFC")
    _set_cell_margins(total_left, top=180, start=180, bottom=180, end=120)
    _set_cell_margins(total_right, top=180, start=120, bottom=180, end=180)
    p1 = total_left.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p1.add_run("Total")
    _set_run_style(run, size=10.5, bold=True, color="64748B")
    p2 = total_right.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p2.add_run(context["total"])
    _set_run_style(run, size=10.5, bold=True)

    _add_body_paragraph(document, "Summary", bold=True, size=12, spacing_after=6)
    summary_table = document.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_width(summary_table, [182])
    summary_cell = summary_table.rows[0].cells[0]
    _shade_cell(summary_cell, "F8FAFC")
    _set_cell_borders(summary_cell)
    _set_cell_margins(summary_cell, top=220, start=220, bottom=220, end=220)
    summary_p = summary_cell.paragraphs[0]
    summary_p.paragraph_format.space_after = Pt(0)
    summary_p.paragraph_format.line_spacing = 1.5
    summary_run = summary_p.add_run(context["summary"])
    _set_run_style(summary_run, size=10.5, color="334155")

    if context.get("notes"):
        _add_body_paragraph(document, "Notes", bold=True, size=12, spacing_after=6)
        _add_body_paragraph(document, context["notes"], size=10.5, color="334155", spacing_after=8)

    if context.get("payment_terms"):
        _add_body_paragraph(document, "Payment Terms", bold=True, size=12, spacing_after=6)
        _add_body_paragraph(document, context["payment_terms"], size=10.5, color="334155", spacing_after=8)

    if any(context.get(key) for key in ("bank_name", "account_name", "account_number", "ifsc", "upi")):
        _add_body_paragraph(document, "Payment Details", bold=True, size=12, spacing_after=6)
        details_table = document.add_table(rows=1, cols=2)
        details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_fixed_width(details_table, [54, 128])
        _clear_table_borders(details_table)
        left, right = details_table.rows[0].cells
        _set_cell_margins(left, top=80, start=0, bottom=80, end=160)
        _set_cell_margins(right, top=80, start=160, bottom=80, end=0)
        for label, value in [
            ("Bank", context.get("bank_name")),
            ("Account Name", context.get("account_name")),
            ("Account No.", context.get("account_number")),
            ("IFSC", context.get("ifsc")),
            ("UPI", context.get("upi")),
        ]:
            if not value:
                continue
            p = left.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run_label = p.add_run(f"{label}  ")
            _set_run_style(run_label, size=10.5, bold=True, color="64748B")
            p2 = right.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            run_value = p2.add_run(str(value))
            _set_run_style(run_value, size=10.5)

    signature = document.add_table(rows=1, cols=2)
    signature.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_width(signature, [50, 132])
    _clear_table_borders(signature)
    seal_cell, auth_cell = signature.rows[0].cells
    _set_cell_margins(seal_cell, top=240, start=0, bottom=40, end=120)
    _set_cell_margins(auth_cell, top=300, start=120, bottom=40, end=0)
    seal_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    auth_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if seal_bytes and context.get("show_seal", True):
        p = seal_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(BytesIO(seal_bytes), width=Mm(28))
    _add_paragraph(auth_cell, context.get("signature_name") or f'Authorized by {context["company_name"]}', size=10.5, spacing_after=2)
    _add_paragraph(auth_cell, context.get("signature_title") or "Authorized Signatory", size=10, color="64748B", spacing_after=6)
    line_p = auth_cell.add_paragraph()
    line_p.paragraph_format.space_after = Pt(0)
    line_run = line_p.add_run("____________________________")
    _set_run_style(line_run, size=10, color="0F172A")

    footer_divider = document.add_paragraph()
    footer_divider.paragraph_format.space_before = Pt(8)
    footer_divider.paragraph_format.space_after = Pt(8)
    footer_run = footer_divider.add_run(" " * 250)
    footer_run.font.underline = True
    footer_run.font.color.rgb = RGBColor.from_string("D7DEE8")

    footer_note = context.get("footer_note")
    _add_body_paragraph(document, f'Thank you for choosing {context["company_name"]}.', size=9, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
    contact_line = context["contact_email"]
    if context["contact_phone"]:
        contact_line = f'{contact_line} | {context["contact_phone"]}'
    _add_body_paragraph(document, contact_line, size=9, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
    if context["address"]:
        _add_body_paragraph(document, context["address"], size=9, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
    if context["gst_number"]:
        _add_body_paragraph(document, f'GST: {context["gst_number"]}', size=9, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
    if footer_note:
        _add_body_paragraph(document, footer_note, size=9, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()


def build_invoice_filename(invoice_id: str, extension: str) -> str:
    safe_id = "".join(char for char in invoice_id if char.isalnum() or char in {"-", "_"})
    return f"{safe_id}.{extension}"
